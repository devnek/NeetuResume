from docx import Document
from docx.shared import Pt

def create_job_docx():
    doc = Document()
    doc.add_heading('Relevant Job Opportunities - Data Analytics & Agriculture', 0)

    jobs = [
        {
            "title": "Manager Demand Planning (Farm PBU)",
            "company": "Farm PBU",
            "description": "Data expert in demand planning meetings, providing data-backed insights and recommendations.",
            "link": "https://www.indeed.com/q-Farm-PBU-jobs.html"
        },
        {
            "title": "Agronomist / Agriculture Consultant (Data-Driven)",
            "company": "Niche Agriculture and Pharmaceuticals Ltd",
            "description": "Precision farming and data-driven agriculture consultant.",
            "link": "https://www.nicheagriculture.com/careers/"
        },
        {
            "title": "Agriculture Expert (Climate)",
            "company": "Wadhwani AI",
            "description": "Using AI and data analysis to address climate-related challenges in agriculture.",
            "link": "https://www.wadhwaniai.org/careers/"
        },
        {
            "title": "Commodity Trader (Agricultural Commodities)",
            "company": "TalentXo",
            "description": "Excel, data analysis, and financial modeling for agricultural commodity portfolios.",
            "link": "https://www.talentxo.com/"
        },
        {
            "title": "Quality Analyst (Food & Agriculture)",
            "company": "S.S. India Food Pvt. Ltd.",
            "description": "Monitoring and testing raw materials and finished goods using data-driven compliance.",
            "link": "https://www.ssindiafood.com/"
        },
        {
            "title": "Senior Data Scientist – Crop Intelligence",
            "company": "MapMyCrop",
            "description": "Satellite AI for real-time intelligence to farmers and insurers.",
            "link": "https://mapmycrop.com/careers/"
        },
        {
            "title": "Strategic Business and Technology Intelligence Consultant",
            "company": "Valeur Fabtex",
            "description": "Market intelligence and strategic consulting for the agriculture domain.",
            "link": "https://valeurfabtex.com/"
        },
        {
            "title": "Precision Agriculture Analyst",
            "company": "PepsiCo India",
            "description": "Supply chain intelligence and crop-level predictive modeling.",
            "link": "https://www.pepsicojobs.com/main/india"
        },
        {
            "title": "Research Analyst (Agricultural Commodity)",
            "company": "IndustryARC",
            "description": "Tracking MSP and global trade policies using agricultural commodity research.",
            "link": "https://www.industryarc.com/careers.php"
        }
    ]

    for job in jobs:
        p = doc.add_paragraph()
        run = p.add_run(job['title'])
        run.bold = True
        run.font.size = Pt(12)
        
        doc.add_paragraph(f"Company: {job['company']}", style='List Bullet')
        doc.add_paragraph(f"Description: {job['description']}", style='List Bullet')
        doc.add_paragraph(f"Link: {job['link']}", style='List Bullet')
        doc.add_paragraph()

    doc.save('Job_Opportunities.docx')
    print("Job_Opportunities.docx created successfully.")

if __name__ == "__main__":
    create_job_docx()
